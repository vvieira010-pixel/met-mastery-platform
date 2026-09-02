from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\Users\vviei\Downloads\met-platform\MET-Mastery-Brand-Identity.docx"

INK = "071527"
TEAL = "01796F"
TEAL_DEEP = "016358"
AQUA = "2DD4BF"
IVORY = "FEFCF5"
MIST = "F4F9FC"
BORDER = "D0E2E8"
PALE_TEAL = "E6F7F4"
SLATE = "3D5A6B"
GREEN = "2D6A4F"
OCHRE = "6B5B3D"
RED = "8B3A3A"
ORANGE = "D97706"
WHITE = "FFFFFF"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_cell_border(cell, color=BORDER, size='6'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'w:' + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)

def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(round(inches * 1440)))
    tc_w.set(qn('w:type'), 'dxa')

def set_run(run, name='DM Sans', size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

def set_spacing(p, before=0, after=6, line=1.25):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def add_text(doc, text='', size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6, line=1.25, font='DM Sans'):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_spacing(p, before, after, line)
    if text:
        set_run(p.add_run(text), font, size, color, bold, italic)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_run(p.add_run(text), size=10.5)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    set_spacing(p, before={1:18, 2:14, 3:10}[level], after={1:10, 2:7, 3:5}[level], line=1.15)
    set_run(p.add_run(text), size={1:16, 2:13, 3:12}[level], color=TEAL if level < 3 else TEAL_DEEP, bold=True)
    return p

def add_table(doc, headers, rows, widths, header_fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i]); set_cell_margins(cell); set_cell_border(cell)
        shade(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; set_spacing(p, after=0, line=1.1)
        set_run(p.add_run(h), size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i]); set_cell_margins(cell); set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]; set_spacing(p, after=0, line=1.15)
            set_run(p.add_run(str(val)), size=9.5, color=INK)
    return table

def add_color_swatch(doc, name, hex_value, role):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    swatch, detail = table.rows[0].cells
    set_cell_width(swatch, 0.65); set_cell_width(detail, 5.85)
    for cell in (swatch, detail):
        set_cell_margins(cell, top=120, bottom=120); set_cell_border(cell)
    shade(swatch, hex_value)
    p = detail.paragraphs[0]; set_spacing(p, after=1, line=1.1)
    set_run(p.add_run(name), size=10.5, color=INK, bold=True)
    p.add_run('  '); set_run(p.add_run('#' + hex_value), name='Space Mono', size=9, color=TEAL_DEEP)
    p = detail.add_paragraph(); set_spacing(p, after=0, line=1.1); set_run(p.add_run(role), size=9.5, color=SLATE)
    return table

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1); section.right_margin = Inches(1)
section.header_distance = Inches(0.35); section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'DM Sans'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'DM Sans'); normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'DM Sans'); normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
for level, size, color in [(1,16,TEAL),(2,13,TEAL),(3,12,TEAL_DEEP)]:
    st = styles[f'Heading {level}']; st.font.name='DM Sans'; st._element.rPr.rFonts.set(qn('w:ascii'),'DM Sans'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'DM Sans'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_spacing(header, after=0, line=1)
set_run(header.add_run('MET MASTERY  /  BRAND IDENTITY'), name='Space Mono', size=8, color=SLATE, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_spacing(footer, after=0, line=1)
set_run(footer.add_run('Clinical English learning workspace  ·  Visual identity guide'), size=8.5, color=SLATE)

# Cover
add_text(doc, 'MET MASTERY', size=11, color=TEAL, bold=True, after=12, font='Space Mono')
p = add_text(doc, 'Visual identity\nfor focused progress', size=31, color=INK, bold=True, after=10, line=1.0)
p.paragraph_format.keep_with_next = True
add_text(doc, 'A practical brand guide for the clinical English learning workspace.', size=14, color=SLATE, after=28, line=1.25)
table = doc.add_table(rows=1, cols=1); table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit=False
cell=table.cell(0,0); set_cell_width(cell,6.5); set_cell_margins(cell,180,180,180,180); set_cell_border(cell, color=TEAL, size='10'); shade(cell, PALE_TEAL)
p=cell.paragraphs[0]; set_spacing(p, after=4, line=1.15); set_run(p.add_run('Brand point of view'), size=11, color=TEAL_DEEP, bold=True)
p=cell.add_paragraph(); set_spacing(p, after=0, line=1.3); set_run(p.add_run('MET Mastery should feel like a calm, evidence-led workspace for healthcare professionals preparing for a high-stakes English exam.'), size=12, color=INK)
add_text(doc, 'Prepared from the implemented MET Mastery interface and design tokens', size=9.5, color=SLATE, italic=True, after=0)

add_heading(doc, '1. Brand essence', 1)
add_text(doc, 'MET Mastery is a teacher-guided English preparation platform built around a clear loop: diagnose, assign targeted practice, review evidence, and improve. The visual identity should make that loop feel trustworthy, focused, and achievable.', after=8)
add_table(doc, ['Attribute', 'Direction'], [
    ('Personality', 'Calm, precise, professional, encouraging'),
    ('Audience', 'Nurses and healthcare professionals preparing for the Michigan English Test'),
    ('Emotional result', '“I know what to do next, and I can see why it matters.”'),
    ('Avoid', 'Childish gamification, noisy edtech gradients, generic medical clichés, hype'),
], [1.55, 4.95])

add_heading(doc, '2. Color system', 1)
add_text(doc, 'Teal is the identity anchor. Deep ink supplies authority. Warm ivory and cool mist create a quiet study environment. Functional colors stay muted so they communicate status without taking over the brand.', after=8)
add_heading(doc, 'Core palette', 2)
for item in [
    ('MET Ink', INK, 'Primary text, dark navigation, authority'),
    ('MET Teal', TEAL, 'Primary actions, links, brand signature'),
    ('MET Teal Deep', TEAL_DEEP, 'Hover, pressed, high-emphasis teal'),
    ('MET Aqua', AQUA, 'Progress, chart accent, active highlight'),
    ('MET Ivory', IVORY, 'Cards, learning surfaces, warm contrast'),
    ('MET Mist', MIST, 'App background and quiet page canvas'),
    ('MET Border', BORDER, 'Dividers, inputs, card boundaries'),
]: add_color_swatch(doc, *item)
add_heading(doc, 'Functional palette', 2)
add_table(doc, ['Meaning', 'Hex', 'Use'], [
    ('Completed / verified', '#'+GREEN, 'Reviewed work, success states'),
    ('Attention needed', '#'+OCHRE, 'Due items, caution, incomplete evidence'),
    ('Error / risk', '#'+RED, 'Blocking errors, risk flags, destructive actions'),
    ('Information', '#'+SLATE, 'Neutral guidance and supporting data'),
    ('Focus / practice', '#'+ORANGE, 'Practice emphasis and focus areas only'),
], [1.55, 1.25, 3.7], header_fill='E8EEF5')

add_heading(doc, '3. Typography', 1)
add_text(doc, 'The type system is intentionally restrained: one humanist UI face for clarity, one editorial serif for rare moments of warmth, and one mono face for measured data.', after=8)
add_table(doc, ['Font', 'Role', 'Guidance'], [
    ('DM Sans', 'Primary UI', 'Use for navigation, forms, dashboards, buttons, body copy, and labels.'),
    ('Cormorant Garamond', 'Editorial display', 'Use sparingly for landing-page statements or campaign moments.'),
    ('Space Mono', 'Measurement', 'Use for timers, scores, IDs, codes, and other system-like values.'),
], [1.35, 1.55, 3.6])
add_heading(doc, 'Recommended scale', 2)
add_table(doc, ['Level', 'Size / weight', 'Use'], [
    ('Display', '48–64px / medium', 'Landing-page statement or major campaign title'),
    ('Page title', '32–40px / bold', 'Dashboard or workspace title'),
    ('Section title', '20–24px / bold', 'Card groups and major sections'),
    ('Body', '16px / regular', 'Explanatory copy and instructions'),
    ('Label', '12–14px / semibold', 'Metadata, controls, status labels'),
    ('Metric', '20–32px / bold mono', 'Scores, counts, timers, percentages'),
], [1.35, 1.55, 3.6])

add_heading(doc, '4. Layout and component language', 1)
add_heading(doc, 'Shape and spacing', 2)
for t in [
    'Use an 8px spacing rhythm. Give headings more space above than below.',
    'Use 8–16px corner radii. Keep the shape soft, but never toy-like.',
    'Use thin pale-blue borders and low-contrast shadows to separate surfaces.',
    'Prefer one clear primary action per section. Hide secondary complexity behind progressive disclosure.',
    'Use cards to organize information, not to decorate every paragraph or create nested containers.',
]: add_bullet(doc, t)
add_heading(doc, 'Buttons and controls', 2)
add_table(doc, ['Pattern', 'Treatment'], [
    ('Primary button', 'MET Teal fill, white text, 8px radius, semibold label'),
    ('Secondary button', 'Transparent or pale teal fill, teal text, thin border'),
    ('Danger action', 'Muted red only when the action is destructive or blocking'),
    ('Status pill', 'Short label, muted background, never the only signal of meaning'),
], [1.65, 4.85])
add_heading(doc, 'Data visualization', 2)
for t in [
    'Use teal for current performance and aqua for active progress.',
    'Use deep ink or slate for targets and previous values.',
    'Use ochre for attention-required states and orange for focus/practice emphasis.',
    'Always show units, visible scales, and a plain-language explanation of what the metric means.',
    'Never rely on color alone; pair color with labels, position, or text.',
]: add_bullet(doc, t)

add_heading(doc, '5. Photography and illustration', 1)
add_text(doc, 'Visuals should make professional preparation feel real and grounded. The audience should recognize themselves as capable adults working toward a meaningful clinical or career goal.', after=8)
add_heading(doc, 'Use', 2)
for t in ['Documentary healthcare and study moments', 'Nurses preparing, communicating, reflecting, or practicing', 'Quiet clinical environments with natural light', 'Subtle details: notes, handover sheets, headphones, study materials, professional conversation']:
    add_bullet(doc, t)
add_heading(doc, 'Avoid', 2)
for t in ['Generic smiling stock-photo teams', 'Overly staged lab-coat imagery', 'Medical cross symbols as decoration', 'Neon technology visuals or exaggerated AI imagery', 'Cartoon classroom graphics and childish gamification']:
    add_bullet(doc, t)

add_heading(doc, '6. Voice and copy', 1)
add_text(doc, 'The voice is direct, calm, and specific. It respects the learner’s professional identity and always connects evidence to a next action.', after=8)
add_table(doc, ['Prefer', 'Avoid'], [
    ('Your next step', 'Level up!'),
    ('Based on evaluated evidence', 'Instant mastery'),
    ('Ready for review', 'Amazing!'),
    ('Focus area', 'Crush the exam!'),
    ('Teacher feedback', 'Guaranteed success'),
], [3.25, 3.25], header_fill='E8EEF5')
add_heading(doc, '7. Canva application rules', 1)
for t in [
    'Use the core palette consistently across presentations, social graphics, worksheets, and teacher materials.',
    'Keep teal as the dominant action and identity color; do not substitute bright blue or purple.',
    'Use warm ivory or cool mist as the main canvas before adding visual decoration.',
    'Limit each design to the core palette plus one functional color.',
    'Use DM Sans for almost all text. Reserve Cormorant Garamond for a deliberate editorial moment.',
    'Make the next action obvious: diagnose, practice, review, or improve.',
]: add_bullet(doc, t)

add_heading(doc, '8. Brand checklist', 1)
add_table(doc, ['Check', 'Pass condition'], [
    ('Color', 'Teal is the identity anchor; functional colors are muted and meaningful.'),
    ('Typography', 'DM Sans leads; display and mono fonts have clear, limited roles.'),
    ('Tone', 'Copy is professional, specific, supportive, and free from hype.'),
    ('Hierarchy', 'One primary action is obvious within a few seconds.'),
    ('Accessibility', 'Color is never the only way to understand status or progress.'),
    ('Imagery', 'Visuals feel documentary, professional, and relevant to healthcare learners.'),
], [1.35, 5.15], header_fill='E8EEF5')
add_text(doc, 'Source note: palette, typography, spacing, and component guidance were extracted from the implemented MET Mastery interface and local design tokens.', size=9, color=SLATE, italic=True, before=16, after=0)

doc.save(OUT)
print(OUT)
